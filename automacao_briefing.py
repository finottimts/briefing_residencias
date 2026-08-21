from pathlib import Path

from pathlib import Path
import re
import unicodedata
import zipfile
from datetime import date
from typing import Iterable

import pandas as pd
from docx import Document
from docx.shared import Pt


# ============================================================
# CONFIGURAÇÃO
# ============================================================

PASTA_BASES = Path(r"")

MODELO_BRIEFING = PASTA_BASES / "Briefing_Estrategico_SGTES_MODELO_14.08.docx"

ARQUIVO_MEDICA_2026 = (
    PASTA_BASES
    / "PANORAMA GERAL - MEDICA_BANCODEDADOS_SIG_ABRIL2026 1.xlsx"
)

ARQUIVO_MULTI_2026 = (
    PASTA_BASES
    / "PANORAMA GERAL - AREA PROFISSIONAL_BANCODEDADOS_SIG_ABRIL2026 1.xlsx"
)

ARQUIVOS_MEDICA = {
    2022: PASTA_BASES / "MEDICA 2022.xlsx",
    2023: PASTA_BASES / "MEDICA 2023.xlsx",
    2024: PASTA_BASES / "MEDICA 2024.xlsx",
    2025: PASTA_BASES / "MEDICA 2025.xlsx",
    2026: ARQUIVO_MEDICA_2026,
}

ARQUIVOS_MULTI = {
    2022: PASTA_BASES / "MULTI 2022.xlsx",
    2023: PASTA_BASES / "MULTI 2023.xlsx",
    2024: PASTA_BASES / "MULTI 2024.xlsx",
    2025: PASTA_BASES / "MULTI 2025.xlsx",
    2026: ARQUIVO_MULTI_2026,
}

PASTA_SAIDA = PASTA_BASES / "Briefings_Gerados"
PASTA_SAIDA.mkdir(parents=True, exist_ok=True)

ANOS = [2022, 2023, 2024, 2025, 2026]

VALOR_BOLSA_COM_PATRONAL = 4927.31

UF_ATUAL = ""


# ============================================================
# TEXTO / NÚMEROS
# ============================================================

def normalizar_texto(valor) -> str:
    if valor is None:
        return ""

    try:
        if pd.isna(valor):
            return ""
    except Exception:
        pass

    texto = str(valor).strip()
    texto = texto.replace("\n", " ").replace("\r", " ").replace("\xa0", " ")

    texto = unicodedata.normalize("NFKD", texto)
    texto = "".join(
        c for c in texto
        if not unicodedata.combining(c)
    )

    return re.sub(r"\s+", " ", texto).upper().strip()


def normalizar_colunas(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    df.columns = [
        re.sub(
            r"\s+",
            " ",
            str(c)
            .replace("\n", " ")
            .replace("\r", " ")
            .replace("\xa0", " ")
            .strip(),
        )
        for c in df.columns
    ]
    return df


def numero(valor) -> float:
    if valor is None:
        return 0.0

    try:
        if pd.isna(valor):
            return 0.0
    except Exception:
        pass

    if isinstance(valor, (int, float)):
        return float(valor)

    texto = str(valor).strip()

    if not texto:
        return 0.0

    texto = texto.replace("R$", "").replace(" ", "")

    if "," in texto and "." in texto:
        texto = texto.replace(".", "").replace(",", ".")
    elif "," in texto:
        texto = texto.replace(",", ".")

    try:
        return float(texto)
    except Exception:
        return 0.0


def formatar_numero(valor) -> str:
    return f"{int(round(numero(valor))):,}".replace(",", ".")


def formatar_moeda(valor) -> str:
    return (
        f"R$ {numero(valor):,.2f}"
        .replace(",", "X")
        .replace(".", ",")
        .replace("X", ".")
    )


def formatar_percentual(valor) -> str:
    if valor is None:
        return "0,0%"

    try:
        if pd.isna(valor):
            return "0,0%"
    except Exception:
        pass

    return f"{float(valor):.1f}%".replace(".", ",")


def variacao(valor_2022, valor_2026):
    v22 = numero(valor_2022)
    v26 = numero(valor_2026)

    absoluta = v26 - v22

    if v22 == 0:
        percentual = 0.0 if v26 == 0 else None
    else:
        percentual = (absoluta / v22) * 100

    return absoluta, percentual


# ============================================================
# COLUNAS
# ============================================================

def localizar_coluna(
    df: pd.DataFrame,
    candidatos: Iterable[str],
    obrigatoria: bool = False,
):
    mapa = {
        normalizar_texto(c): c
        for c in df.columns
    }

    candidatos = list(candidatos)

    for candidato in candidatos:
        chave = normalizar_texto(candidato)
        if chave in mapa:
            return mapa[chave]

    for coluna in df.columns:
        n_coluna = normalizar_texto(coluna)

        for candidato in candidatos:
            n_candidato = normalizar_texto(candidato)

            if (
                n_candidato
                and (
                    n_candidato in n_coluna
                    or n_coluna in n_candidato
                )
            ):
                return coluna

    if obrigatoria:
        raise KeyError(
            "Coluna obrigatória não encontrada.\n"
            f"Procuradas: {candidatos}\n"
            f"Disponíveis: {list(df.columns)}"
        )

    return None


# ============================================================
# LEITURA DAS BASES
# ============================================================

def ler_excel_com_header(arquivo: Path) -> pd.DataFrame:
    ultimo_erro = None

    for header in range(10):
        try:
            df = pd.read_excel(arquivo, header=header)
            df = normalizar_colunas(df)

            colunas = {
                normalizar_texto(c)
                for c in df.columns
            }

            if (
                "UF" in colunas
                and (
                    "BOLSAS FINANCIADAS" in colunas
                    or "BOLSA FINANCIADA" in colunas
                    or "BOLSAS" in colunas
                )
            ):
                return df

        except Exception as exc:
            ultimo_erro = exc

    if ultimo_erro:
        raise ultimo_erro

    return normalizar_colunas(pd.read_excel(arquivo))


def preparar_base(df: pd.DataFrame, ano: int, tipo: str) -> pd.DataFrame:
    df = normalizar_colunas(df)

    col_uf = localizar_coluna(
        df,
        ["uf"],
        obrigatoria=True,
    )

    col_municipio = localizar_coluna(
        df,
        [
            "no_municipio",
            "municipio",
            "município",
            "município de oferta",
        ],
    )

    col_programa = localizar_coluna(
        df,
        [
            "programa",
            "nome do programa",
        ],
        obrigatoria=True,
    )

    col_bolsas = localizar_coluna(
        df,
        [
            "bolsas financiadas",
            "bolsa financiada",
            "bolsas",
            "quantidade de bolsas",
        ],
        obrigatoria=True,
    )

    col_instituicao = localizar_coluna(
        df,
        [
            "no_razao_social",
            "razao social",
            "razão social",
            "no_fantasia",
            "instituição",
            "instituicao",
            "nome da instituição",
        ],
    )

    col_codigo = localizar_coluna(
        df,
        [
            "co_seq_programa",
            "co_programa",
            "codigo programa",
            "código programa",
        ],
    )

    df["_UF_"] = df[col_uf].map(normalizar_texto)

    if col_municipio:
        df["_MUNICIPIO_"] = df[col_municipio].map(normalizar_texto)
    else:
        df["_MUNICIPIO_"] = ""

    df["_PROGRAMA_"] = (
        df[col_programa]
        .fillna("")
        .astype(str)
        .str.strip()
    )

    df["_BOLSAS_"] = df[col_bolsas].map(numero)

    # Somente programas com bolsas financiadas.
    df = df[df["_BOLSAS_"] > 0].copy()

    if col_instituicao:
        df["_INSTITUICAO_"] = (
            df[col_instituicao]
            .fillna("")
            .astype(str)
            .str.strip()
        )
    else:
        df["_INSTITUICAO_"] = ""

    if col_codigo:
        df["_COD_PROGRAMA_"] = (
            df[col_codigo]
            .fillna("")
            .astype(str)
            .str.strip()
        )
    else:
        df["_COD_PROGRAMA_"] = ""

    df["_ANO_"] = int(ano)
    df["_TIPO_"] = tipo

    ids = []

    for _, row in df.iterrows():
        codigo = str(row["_COD_PROGRAMA_"]).strip()
        programa = normalizar_texto(row["_PROGRAMA_"])
        instituicao = normalizar_texto(row["_INSTITUICAO_"])

        if codigo and codigo.upper() not in {"NAN", "NONE"}:
            identificador = f"COD:{codigo}"
        elif programa:
            identificador = (
                f"INST:{instituicao}|PROG:{programa}"
            )
        else:
            identificador = ""

        ids.append(identificador)

    df["_ID_PROGRAMA_"] = ids

    return df


def carregar_bases():
    bases_medica = {}
    bases_multi = {}

    print("\nCarregando bases históricas...")

    for ano in ANOS:
        arq_med = ARQUIVOS_MEDICA[ano]
        arq_multi = ARQUIVOS_MULTI[ano]

        if not arq_med.exists():
            raise FileNotFoundError(
                f"Arquivo de Residência Médica não encontrado:\n{arq_med}"
            )

        if not arq_multi.exists():
            raise FileNotFoundError(
                f"Arquivo de Área Profissional não encontrado:\n{arq_multi}"
            )

        print(f"  Médica {ano}...")
        bases_medica[ano] = preparar_base(
            ler_excel_com_header(arq_med),
            ano,
            "medica",
        )

        print(f"  Multi {ano}...")
        bases_multi[ano] = preparar_base(
            ler_excel_com_header(arq_multi),
            ano,
            "multi",
        )

        print(
            f"    ✓ {ano}: "
            f"{len(bases_medica[ano]):,} médica | "
            f"{len(bases_multi[ano]):,} multi"
        )

    return bases_medica, bases_multi


# ============================================================
# FILTROS
# ============================================================

def filtrar(
    df: pd.DataFrame,
    uf: str | None = None,
    municipio: str | None = None,
    instituicoes: list[str] | None = None,
) -> pd.DataFrame:

    dados = df.copy()

    if uf:
        dados = dados[
            dados["_UF_"] == normalizar_texto(uf)
        ]

    if municipio:
        dados = dados[
            dados["_MUNICIPIO_"] == normalizar_texto(municipio)
        ]

    if instituicoes:
        selecionadas = {
            normalizar_texto(x)
            for x in instituicoes
        }

        dados = dados[
            dados["_INSTITUICAO_"]
            .map(normalizar_texto)
            .isin(selecionadas)
        ]

    return dados


def listar_ufs(bases_medica, bases_multi):
    ufs = set()

    for df in bases_medica.values():
        ufs.update(df["_UF_"].dropna().unique())

    for df in bases_multi.values():
        ufs.update(df["_UF_"].dropna().unique())

    return sorted(x for x in ufs if x)


def listar_municipios(
    bases_medica,
    bases_multi,
    uf,
):
    municipios = set()
    uf_norm = normalizar_texto(uf)

    for df in list(bases_medica.values()) + list(bases_multi.values()):
        dados = df[df["_UF_"] == uf_norm]
        municipios.update(
            dados["_MUNICIPIO_"].dropna().unique()
        )

    return sorted(x for x in municipios if x)


def listar_instituicoes(
    bases_medica,
    bases_multi,
    uf,
    municipio=None,
):
    instituicoes = set()

    for df in (
        bases_medica[2026],
        bases_multi[2026],
    ):
        dados = filtrar(df, uf, municipio)

        instituicoes.update(
            x
            for x in dados["_INSTITUICAO_"].dropna().astype(str)
            if x.strip()
        )

    return sorted(
        instituicoes,
        key=normalizar_texto,
    )


# ============================================================
# INDICADORES
# ============================================================

def contar_programas(df: pd.DataFrame) -> int:
    if df.empty:
        return 0

    ids = (
        df["_ID_PROGRAMA_"]
        .replace("", pd.NA)
        .dropna()
    )

    return int(ids.nunique())


def contar_bolsas(df: pd.DataFrame) -> int:
    if df.empty:
        return 0

    return int(round(df["_BOLSAS_"].sum()))


def financeiro(bolsas) -> float:
    return numero(bolsas) * VALOR_BOLSA_COM_PATRONAL


def calcular_indicadores_ano(
    df_medica,
    df_multi,
    uf=None,
    municipio=None,
    instituicoes=None,
):
    med = filtrar(
        df_medica,
        uf,
        municipio,
        instituicoes,
    )

    multi = filtrar(
        df_multi,
        uf,
        municipio,
        instituicoes,
    )

    programas_medica = contar_programas(med)
    programas_multi = contar_programas(multi)

    bolsas_medica = contar_bolsas(med)
    bolsas_multi = contar_bolsas(multi)

    return {
        "programas_medica": programas_medica,
        "programas_multi": programas_multi,
        "programas_total": programas_medica + programas_multi,
        "bolsas_medica": bolsas_medica,
        "bolsas_multi": bolsas_multi,
        "bolsas_total": bolsas_medica + bolsas_multi,
        "financeiro_medica": financeiro(bolsas_medica),
        "financeiro_multi": financeiro(bolsas_multi),
        "financeiro_total": financeiro(
            bolsas_medica + bolsas_multi
        ),
        "_medica": med,
        "_multi": multi,
    }


def calcular_serie(
    bases_medica,
    bases_multi,
    uf=None,
    municipio=None,
    instituicoes=None,
):
    serie = {}

    for ano in ANOS:
        indicadores = calcular_indicadores_ano(
            bases_medica[ano],
            bases_multi[ano],
            uf,
            municipio,
            instituicoes,
        )

        serie[ano] = {
            chave: indicadores[chave]
            for chave in [
                "programas_medica",
                "programas_multi",
                "programas_total",
                "bolsas_medica",
                "bolsas_multi",
                "bolsas_total",
                "financeiro_medica",
                "financeiro_multi",
                "financeiro_total",
            ]
        }

    return serie


# ============================================================
# TOP 5 INSTITUIÇÕES POR PROGRAMAS
# ============================================================

def calcular_top5_instituicoes_programas(
    df_medica,
    df_multi,
    uf,
    municipio,
):
    """
    TOP 5 por quantidade de PROGRAMAS.

    Os dados são exclusivamente do município selecionado.
    A quantidade de bolsas aparece apenas como informação
    complementar.
    """

    if not municipio:
        vazio = pd.DataFrame(
            columns=[
                "Instituição",
                "Programas",
                "Bolsas financiadas",
            ]
        )
        return vazio.copy(), vazio.copy()

    medica = filtrar(
        df_medica,
        uf=uf,
        municipio=municipio,
    ).copy()

    multi = filtrar(
        df_multi,
        uf=uf,
        municipio=municipio,
    ).copy()

    def calcular(df):
        if df.empty:
            return pd.DataFrame(
                columns=[
                    "Instituição",
                    "Programas",
                    "Bolsas financiadas",
                ]
            )

        df["_INSTITUICAO_NORM_"] = (
            df["_INSTITUICAO_"]
            .fillna("")
            .astype(str)
            .map(normalizar_texto)
        )

        df = df[
            df["_INSTITUICAO_NORM_"] != ""
        ].copy()

        resultados = []

        for chave, grupo in df.groupby(
            "_INSTITUICAO_NORM_",
            sort=False,
        ):
            nome = (
                grupo["_INSTITUICAO_"]
                .dropna()
                .astype(str)
                .str.strip()
            )

            nome = (
                nome.iloc[0]
                if not nome.empty
                else chave
            )

            programas = (
                grupo["_ID_PROGRAMA_"]
                .replace("", pd.NA)
                .dropna()
                .nunique()
            )

            bolsas = grupo["_BOLSAS_"].sum()

            if programas > 0:
                resultados.append(
                    {
                        "Instituição": nome,
                        "Programas": int(programas),
                        "Bolsas financiadas": int(round(bolsas)),
                    }
                )

        resultado = pd.DataFrame(resultados)

        if resultado.empty:
            return pd.DataFrame(
                columns=[
                    "Instituição",
                    "Programas",
                    "Bolsas financiadas",
                ]
            )

        return (
            resultado
            .sort_values(
                by=[
                    "Programas",
                    "Bolsas financiadas",
                    "Instituição",
                ],
                ascending=[False, False, True],
            )
            .head(5)
            .reset_index(drop=True)
        )

    return calcular(medica), calcular(multi)


# ============================================================
# WORD - AUXILIARES
# ============================================================

def escrever_celula(
    celula,
    valor,
    tamanho=10,
    negrito=False,
):
    celula.text = ""

    p = celula.paragraphs[0]
    run = p.add_run(str(valor))
    run.font.size = Pt(tamanho)
    run.bold = negrito


def iterar_paragrafos_documento(doc):
    """
    Retorna parágrafos do corpo do documento e também das células
    das tabelas.
    """

    for p in doc.paragraphs:
        yield p

    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for p in celula.paragraphs:
                    yield p


def substituir_no_texto_documento(
    doc,
    substituicoes,
):
    """
    Substituição somente textual.
    Não trata [nº].
    """

    for p in iterar_paragrafos_documento(doc):
        texto = p.text
        novo = texto

        for antigo, valor in substituicoes.items():
            novo = novo.replace(
                antigo,
                str(valor),
            )

        if novo != texto:
            p.text = novo


def inserir_tabela_apos_paragrafo(
    doc,
    texto_referencia,
    headers,
    linhas,
):
    """
    Insere uma tabela REAL imediatamente depois do parágrafo
    que contém o texto_referencia.

    O marcador é removido para não permanecer no documento.
    """

    referencia = normalizar_texto(texto_referencia)
    paragrafo_alvo = None

    for p in iterar_paragrafos_documento(doc):
        if referencia in normalizar_texto(p.text):
            paragrafo_alvo = p
            break

    if paragrafo_alvo is None:
        raise RuntimeError(
            "Marcador da tabela não encontrado no modelo:\n"
            f"{texto_referencia}"
        )

    # Mantém o parágrafo como título, mas remove textos de marcador
    # quando o título é apenas o marcador.
    texto_atual = paragrafo_alvo.text.strip()

    if (
        normalizar_texto(texto_atual)
        == referencia
    ):
        paragrafo_alvo.text = ""

    tabela = doc.add_table(
        rows=1,
        cols=len(headers),
    )

    tabela.style = "Table Grid"

    for coluna, valor in enumerate(headers):
        escrever_celula(
            tabela.cell(0, coluna),
            valor,
            tamanho=9,
            negrito=True,
        )

    for linha_dados in linhas:
        cells = tabela.add_row().cells

        for coluna, valor in enumerate(linha_dados):
            escrever_celula(
                cells[coluna],
                valor,
                tamanho=9,
            )

    # Move a tabela para depois do parágrafo.
    paragrafo_alvo._p.addnext(tabela._tbl)

    return tabela


# ============================================================
# GRANDES NÚMEROS - UF
# ============================================================

def preencher_grandes_numeros(
    doc,
    indicadores,
    uf,
):
    """
    Primeiro quadro sempre representa a UF inteira.

    Inclui os valores de Área Profissional tanto em bolsas
    quanto em financeiro.
    """

    if len(doc.tables) < 2:
        raise RuntimeError(
            "O modelo não possui a tabela de grandes números esperada."
        )

    tabela = doc.tables[1]

    if len(tabela.rows) != 1 or len(tabela.rows[0].cells) != 2:
        raise RuntimeError(
            "A estrutura da tabela de grandes números mudou. "
            "Esperado: 1 linha x 2 células."
        )

    esquerda = tabela.rows[0].cells[0]
    direita = tabela.rows[0].cells[1]

    esquerda.text = (
        f"Residências em saúde {normalizar_texto(uf)}\n"
        "(BOLSAS DE RESIDÊNCIA FINANCIADAS PELO MS)\n"
        f"{formatar_numero(indicadores['programas_total'])}\n"
        "Residentes em saúde com bolsa financiada pelo MS\n"
        f"{formatar_numero(indicadores['bolsas_total'])}\n"
        "Em Residência Médica\n"
        f"{formatar_numero(indicadores['bolsas_medica'])}\n"
        "Em Área Profissional\n"
        f"{formatar_numero(indicadores['bolsas_multi'])}"
    )

    direita.text = (
        "Valor financeiro residências ($)\n"
        "(BOLSAS DE RESIDÊNCIA FINANCIADAS PELO MS)\n"
        f"{formatar_moeda(indicadores['financeiro_total'])}\n"
        "Residentes no Estado em saúde com bolsa financiada pelo MS\n"
        f"{formatar_moeda(indicadores['financeiro_total'])}\n"
        "Em Residência Médica\n"
        f"{formatar_moeda(indicadores['financeiro_medica'])}\n"
        "Em Área Profissional\n"
        f"{formatar_moeda(indicadores['financeiro_multi'])}"
    )


# ============================================================
# VARIAÇÃO
# ============================================================

def preencher_variacoes_residencias(
    doc,
    serie_brasil,
    serie_estado,
):
    if len(doc.tables) < 4:
        raise RuntimeError(
            "O modelo não possui as tabelas de variação esperadas."
        )

    v_abs_brasil, v_pct_brasil = variacao(
        serie_brasil[2022]["bolsas_total"],
        serie_brasil[2026]["bolsas_total"],
    )

    v_abs_estado, v_pct_estado = variacao(
        serie_estado[2022]["bolsas_total"],
        serie_estado[2026]["bolsas_total"],
    )

    doc.tables[2].cell(0, 0).text = (
        f"Residências: {formatar_numero(v_abs_brasil)} "
        f"({formatar_percentual(v_pct_brasil)})"
    )

    doc.tables[3].cell(0, 0).text = (
        f"Residências: {formatar_numero(v_abs_estado)} "
        f"({formatar_percentual(v_pct_estado)})"
    )


# ============================================================
# COMPARATIVO BRASIL X UF
# ============================================================

def preencher_comparativo(
    doc,
    brasil,
    estado,
    uf,
):
    if len(doc.tables) < 5:
        raise RuntimeError(
            "Tabela comparativa não encontrada."
        )

    tabela = doc.tables[4]

    if len(tabela.rows) < 10 or len(tabela.columns) < 4:
        raise RuntimeError(
            "A estrutura da tabela comparativa mudou."
        )

    linhas = {
        1: ("programas_medica", False),
        2: ("programas_multi", False),
        3: ("programas_total", False),
        4: ("bolsas_medica", False),
        5: ("bolsas_multi", False),
        6: ("bolsas_total", False),
        7: ("financeiro_medica", True),
        8: ("financeiro_multi", True),
        9: ("financeiro_total", True),
    }

    for indice_linha, (chave, eh_moeda) in linhas.items():

        valor_brasil = brasil[chave]
        valor_estado = estado[chave]

        if eh_moeda:
            texto_brasil = formatar_moeda(valor_brasil)
            texto_estado = formatar_moeda(valor_estado)
        else:
            texto_brasil = formatar_numero(valor_brasil)
            texto_estado = formatar_numero(valor_estado)

        if numero(valor_brasil) == 0:
            participacao = 0.0
        else:
            participacao = (
                numero(valor_estado)
                / numero(valor_brasil)
                * 100
            )

        escrever_celula(
            tabela.cell(indice_linha, 1),
            texto_brasil,
        )

        escrever_celula(
            tabela.cell(indice_linha, 2),
            texto_estado,
        )

        escrever_celula(
            tabela.cell(indice_linha, 3),
            formatar_percentual(participacao),
        )

    tabela.cell(0, 2).text = (
        f"Estado ({normalizar_texto(uf)})"
    )


# ============================================================
# SÉRIE HISTÓRICA - UF
# ============================================================

def preencher_serie_historica(
    doc,
    serie_brasil,
    serie_estado,
):
    if len(doc.tables) < 6:
        raise RuntimeError(
            "Tabela de série histórica não encontrada."
        )

    tabela = doc.tables[5]

    if len(tabela.rows) < 10 or len(tabela.columns) < 8:
        raise RuntimeError(
            "A estrutura da série histórica mudou."
        )

    linhas = {
        1: ("programas_medica", False),
        2: ("programas_multi", False),
        3: ("programas_total", False),
        4: ("bolsas_medica", False),
        5: ("bolsas_multi", False),
        6: ("bolsas_total", False),
        7: ("financeiro_medica", True),
        8: ("financeiro_multi", True),
        9: ("financeiro_total", True),
    }

    for indice_linha, (chave, eh_moeda) in linhas.items():

        for posicao, ano in enumerate(ANOS, start=1):

            valor = serie_estado[ano][chave]

            texto = (
                formatar_moeda(valor)
                if eh_moeda
                else formatar_numero(valor)
            )

            escrever_celula(
                tabela.cell(indice_linha, posicao),
                texto,
            )

        valor_2022 = serie_estado[2022][chave]
        valor_2026 = serie_estado[2026][chave]

        v_abs, v_pct = variacao(
            valor_2022,
            valor_2026,
        )

        texto_abs = (
            formatar_moeda(v_abs)
            if eh_moeda
            else formatar_numero(v_abs)
        )

        escrever_celula(
            tabela.cell(indice_linha, 6),
            texto_abs,
        )

        escrever_celula(
            tabela.cell(indice_linha, 7),
            formatar_percentual(v_pct),
        )


# ============================================================
# GRÁFICOS
# ============================================================

def _substituir_valores_chart_xml(
    xml: bytes,
    valores: list,
) -> bytes:

    texto = xml.decode("utf-8")

    padrao_cache = re.compile(
        r"(<c:numCache>.*?</c:numCache>)",
        flags=re.DOTALL,
    )

    match = padrao_cache.search(texto)

    if not match:
        return xml

    cache = match.group(1)

    def trocar_pt(match_pt):
        bloco = match_pt.group(0)

        idx = re.search(
            r'<c:pt idx="(\d+)">',
            bloco,
        )

        if not idx:
            return bloco

        pos = int(idx.group(1))

        if pos >= len(valores):
            return bloco

        bloco = re.sub(
            r"<c:v>.*?</c:v>",
            f"<c:v>{numero(valores[pos])}</c:v>",
            bloco,
            count=1,
        )

        return bloco

    novo_cache = re.sub(
        r'<c:pt idx="\d+">.*?</c:pt>',
        trocar_pt,
        cache,
        flags=re.DOTALL,
    )

    texto = (
        texto[:match.start(1)]
        + novo_cache
        + texto[match.end(1):]
    )

    return texto.encode("utf-8")


def atualizar_graficos(
    caminho_docx: Path,
    serie_brasil,
    serie_estado,
):
    valores_brasil = [
        serie_brasil[ano]["bolsas_total"]
        for ano in ANOS
    ]

    valores_estado = [
        serie_estado[ano]["bolsas_total"]
        for ano in ANOS
    ]

    temporario = caminho_docx.with_suffix(".tmp.docx")

    with zipfile.ZipFile(
        caminho_docx,
        "r",
    ) as zin, zipfile.ZipFile(
        temporario,
        "w",
        compression=zipfile.ZIP_DEFLATED,
    ) as zout:

        for item in zin.infolist():
            dados = zin.read(item.filename)

            if item.filename == "word/charts/chart1.xml":
                dados = _substituir_valores_chart_xml(
                    dados,
                    valores_brasil,
                )

            elif item.filename == "word/charts/chart2.xml":
                dados = _substituir_valores_chart_xml(
                    dados,
                    valores_estado,
                )

            zout.writestr(item, dados)

    temporario.replace(caminho_docx)


# ============================================================
# SEÇÃO 3.2 - MUNICÍPIO
# ============================================================

def obter_dados_municipio_2026(
    bases_medica,
    bases_multi,
    uf,
    municipio,
):
    """
    Dados exclusivos do município selecionado.
    """

    return calcular_indicadores_ano(
        bases_medica[2026],
        bases_multi[2026],
        uf=uf,
        municipio=municipio,
        instituicoes=None,
    )


def preencher_texto_secao_32(
    doc,
    uf,
    municipio,
    indicadores_municipio,
):
    """
    Preenche especificamente a seção 3.2 com os dados
    do município.

    IMPORTANTE:
    Os indicadores gerais continuam sendo da UF.
    """

    if not municipio:
        return

    municipio_exibicao = str(municipio).strip()
    uf_exibicao = normalizar_texto(uf)

    texto_secao = (
        f"3.2 INFORMAÇÕES DETALHADAS DOS PROGRAMAS DE "
        f"RESIDÊNCIA EM {municipio_exibicao}-{uf_exibicao}:"
    )

    texto_corpo = (
        f"No município de {municipio_exibicao}, foram identificados "
        f"{formatar_numero(indicadores_municipio['programas_total'])} "
        f"programas de residência em saúde, distribuídos entre "
        f"{formatar_numero(indicadores_municipio['programas_medica'])} "
        f"programas médicos e "
        f"{formatar_numero(indicadores_municipio['programas_multi'])} "
        f"programas em área profissional da saúde. "
        f"O total de bolsas financiadas no município é de "
        f"{formatar_numero(indicadores_municipio['bolsas_total'])}, "
        f"distribuídas entre "
        f"{formatar_numero(indicadores_municipio['bolsas_medica'])} "
        f"bolsas em residência médica e "
        f"{formatar_numero(indicadores_municipio['bolsas_multi'])} "
        f"bolsas em residência em área profissional da saúde."
    )

    texto_medica = (
        "No município, os programas de Residência Médica com bolsas "
        "financiadas pelo Ministério da Saúde estão distribuídos "
        "nas seguintes instituições."
    )

    texto_multi = (
        "No município, os programas de Residência Área Profissional da Saúde"
        "com bolsas financiadas pelo Ministério da Saúde estão "
        "distribuídos nas seguintes instituições."
    )

    substituicoes = {
        "[3.2 INFORMAÇÕES DETALHADAS DOS PROGRAMAS DE RESIDÊNCIA EM [município]-[UF]]":
            texto_secao,
        "[X] programas de residência em saúde":
            f"{formatar_numero(indicadores_municipio['programas_total'])} programas de residência em saúde",
        "[X] programas médicos":
            f"{formatar_numero(indicadores_municipio['programas_medica'])} programas médicos",
        "[X] programas em área profissional da saúde":
            f"{formatar_numero(indicadores_municipio['programas_multi'])} programas em área profissional da saúde",
        "[X] bolsas financiadas":
            f"{formatar_numero(indicadores_municipio['bolsas_total'])} bolsas financiadas",
        "[X] bolsas em residência médica":
            f"{formatar_numero(indicadores_municipio['bolsas_medica'])} bolsas em residência médica",
        "[X] bolsas em residência em área profissional da saúde":
            f"{formatar_numero(indicadores_municipio['bolsas_multi'])} bolsas em residência em área profissional da saúde",
    }

    substituir_no_texto_documento(
        doc,
        substituicoes,
    )

    # --------------------------------------------------------
    # Procura e substitui o título da seção 3.2
    # --------------------------------------------------------

    for p in iterar_paragrafos_documento(doc):
        normalizado = normalizar_texto(p.text)

        if (
            "3.2" in normalizado
            and "INFORMACOES DETALHADAS" in normalizado
            and "RESIDENCIA" in normalizado
        ):
            p.text = texto_secao

        elif (
            "NO MUNICIPIO DE" in normalizado
            and "FORAM IDENTIFICADOS" in normalizado
            and "PROGRAMAS DE RESIDENCIA" in normalizado
        ):
            p.text = texto_corpo

        elif (
            "NO MUNICIPIO" in normalizado
            and "RESIDENCIA MEDICA" in normalizado
            and "INSTITUICOES" in normalizado
        ):
            p.text = texto_medica

        elif (
            "RESIDENCIA EM AREA PROFISSIONAL" in normalizado
            and "INSTITUICOES" in normalizado
        ):
            p.text = texto_multi


# ============================================================
# TABELAS TOP 5
# ============================================================

def encontrar_marcador_top5(doc, tipo):
    """
    Localiza os marcadores da tabela no modelo.

    Aceita variações antigas e novas do texto.
    """

    if tipo == "medica":
        candidatos = [
            "Tabela Instituições com maior número de programas instituições medicas",
            "Tabela Instituições com maior número de programas instituições médica",
            "Tabela Instituições com maior número de programas instituições medica",
            "5 instituições",
            "5 insitituições",
        ]
    else:
        candidatos = [
            "Tabela Instituições com maior número de programas instituições multi",
            "Tabela Instituições com maior número de programas instituições área profissional",
            "Tabela Instituições com maior número de programas instituições multiprofissional",
            "5 instituições",
            "5 insitituições",
        ]

    candidatos_norm = [
        normalizar_texto(x)
        for x in candidatos
    ]

    # Primeiro tenta os marcadores específicos.
    for p in iterar_paragrafos_documento(doc):
        texto = normalizar_texto(p.text)

        for candidato in candidatos_norm[:4]:
            if candidato and candidato in texto:
                return p

    return None


def inserir_tabela_top5(
    doc,
    tipo,
    dados,
):
    """
    Insere a tabela real de Top 5.

    Ranking = número de programas.
    """

    paragrafo = encontrar_marcador_top5(
        doc,
        tipo,
    )

    if paragrafo is None:
        raise RuntimeError(
            "Não encontrei o local reservado para a tabela "
            f"Top 5 de {tipo} no modelo."
        )

    if tipo == "medica":
        titulo = (
            "Instituições com maior número de programas "
            "de Residência Médica"
        )
    else:
        titulo = (
            "Instituições com maior número de programas "
            "de Residência em Área Profissional da Saúde"
        )

    # O marcador vira o título correto.
    paragrafo.text = titulo

    linhas = []

    if dados.empty:
        linhas.append(
            [
                "-",
                "Nenhuma instituição encontrada",
                "0",
                "0",
            ]
        )
    else:
        for posicao, (_, linha) in enumerate(
            dados.iterrows(),
            start=1,
        ):
            linhas.append(
                [
                    str(posicao),
                    str(linha["Instituição"]),
                    formatar_numero(linha["Programas"]),
                    formatar_numero(
                        linha["Bolsas financiadas"]
                    ),
                ]
            )

    tabela = doc.add_table(
        rows=1,
        cols=4,
    )

    tabela.style = "Table Grid"

    headers = [
        "Posição",
        "Instituição",
        "Programas",
        "Bolsas financiadas",
    ]

    for i, header in enumerate(headers):
        escrever_celula(
            tabela.cell(0, i),
            header,
            tamanho=9,
            negrito=True,
        )

    for linha in linhas:
        cells = tabela.add_row().cells

        for i, valor in enumerate(linha):
            escrever_celula(
                cells[i],
                valor,
                tamanho=9,
            )

    paragrafo._p.addnext(tabela._tbl)

    return tabela


def preencher_tabelas_top5(
    doc,
    bases_medica,
    bases_multi,
    uf,
    municipio,
):
    """
    Cria as duas tabelas Top 5 usando somente o município.

    Não considera a lista de instituições escolhida no terminal.
    """

    if not municipio:
        print(
            "\nATENÇÃO: nenhum município foi selecionado."
        )
        print(
            "As tabelas Top 5 do município não serão geradas."
        )
        return

    top_medica, top_multi = (
        calcular_top5_instituicoes_programas(
            bases_medica[2026],
            bases_multi[2026],
            uf,
            municipio,
        )
    )

    inserir_tabela_top5(
        doc,
        "medica",
        top_medica,
    )

    inserir_tabela_top5(
        doc,
        "multi",
        top_multi,
    )


# ============================================================
# TEXTOS GERAIS
# ============================================================

def nome_mes(mes):
    nomes = {
        1: "janeiro",
        2: "fevereiro",
        3: "março",
        4: "abril",
        5: "maio",
        6: "junho",
        7: "julho",
        8: "agosto",
        9: "setembro",
        10: "outubro",
        11: "novembro",
        12: "dezembro",
    }

    return nomes[mes]


def preencher_textos(
    doc,
    uf,
    municipio,
    indicadores_uf,
):
    hoje = date.today()

    uf_exibicao = normalizar_texto(uf)

    contexto = (
        f"No estado de {uf_exibicao}, foram identificados "
        f"{formatar_numero(indicadores_uf['programas_total'])} "
        f"programas de residência em saúde, com "
        f"{formatar_numero(indicadores_uf['bolsas_total'])} "
        f"bolsas financiadas pelo Ministério da Saúde."
    )

    substituicoes = {
        "[UF]": uf_exibicao,
        "[uf]": uf_exibicao,
        "[dia]": f"{hoje.day:02d}",
        "[mês]": nome_mes(8),
        "[mês]": nome_mes(8),
        "[nome do mês]": nome_mes(8),
        "[Nome do mês]": nome_mes(8),
        "[ano]": "2026",
        "[inserir texto]": contexto,
        "[Inserir texto]": contexto,
    }

    # Atenção: não substituir [município] aqui.
    # O município pertence à seção 3.2.
    substituir_no_texto_documento(
        doc,
        substituicoes,
    )

    for p in doc.paragraphs:
        texto = p.text
        normalizado = normalizar_texto(texto)

        if "EVOLUÇÃO DOS INDICADORES BRASIL" in normalizado:
            p.text = (
                "EVOLUÇÃO DOS INDICADORES BRASIL - (2022 – 2026)"
            )

        elif "EVOLUÇÃO DOS INDICADORES NO ESTADO" in normalizado:
            p.text = (
                f"EVOLUÇÃO DOS INDICADORES NO ESTADO "
                f"{uf_exibicao} - (2022 – 2026)"
            )

        elif "3. RESIDÊNCIAS MÉDICAS E EM SAÚDE" in normalizado:
            p.text = (
                "3. RESIDÊNCIAS MÉDICAS E EM SAÚDE "
                "(INDICADORES SAGE & SÉRIE HISTÓRICA 2022-2026)"
            )

        elif normalizado.startswith("VARIAÇÃO"):
            p.text = (
                "Variação 2022–2026 (absoluta e percentual):"
            )

    for section in doc.sections:
        for p in section.footer.paragraphs:
            if "[mês]" in p.text or "[ano]" in p.text:
                p.text = (
                    "Briefing SGTES | Elaboração: SGTES/MS | "
                    f"Data: {hoje.day:02d}/{nome_mes(8)}/2026"
                )


# ============================================================
# VALIDAÇÃO
# ============================================================

def encontrar_placeholders(doc):
    encontrados = []

    for p in iterar_paragrafos_documento(doc):
        if re.search(
            r"\[Nº\]",
            p.text,
            flags=re.IGNORECASE,
        ):
            encontrados.append(
                ("PARÁGRAFO", p.text)
            )

    return encontrados


def validar_valores(
    indicadores_estado,
    serie_estado,
):
    erros = []

    if indicadores_estado["programas_total"] != (
        indicadores_estado["programas_medica"]
        + indicadores_estado["programas_multi"]
    ):
        erros.append(
            "Total de programas não corresponde à soma Médica + Multi."
        )

    if indicadores_estado["bolsas_total"] != (
        indicadores_estado["bolsas_medica"]
        + indicadores_estado["bolsas_multi"]
    ):
        erros.append(
            "Total de bolsas não corresponde à soma Médica + Multi."
        )

    esperado = financeiro(
        indicadores_estado["bolsas_total"]
    )

    if round(
        indicadores_estado["financeiro_total"],
        2,
    ) != round(
        esperado,
        2,
    ):
        erros.append(
            "Financeiro total não corresponde às bolsas x valor unitário."
        )

    for ano in ANOS:
        if serie_estado[ano]["bolsas_total"] != (
            serie_estado[ano]["bolsas_medica"]
            + serie_estado[ano]["bolsas_multi"]
        ):
            erros.append(
                f"Série {ano}: total de bolsas inconsistente."
            )

    if erros:
        raise ValueError(
            "\n".join(
                f"- {erro}"
                for erro in erros
            )
        )


# ============================================================
# GERAÇÃO
# ============================================================

def gerar_briefing(
    bases_medica,
    bases_multi,
    uf,
    municipio=None,
    instituicoes=None,
):
    global UF_ATUAL

    UF_ATUAL = normalizar_texto(uf)

    # ========================================================
    # IMPORTANTE:
    # INDICADORES GERAIS = UF INTEIRA
    # ========================================================

    indicadores_uf = calcular_indicadores_ano(
        bases_medica[2026],
        bases_multi[2026],
        uf=uf,
        municipio=None,
        instituicoes=None,
    )

    # ========================================================
    # BRASIL
    # ========================================================

    indicadores_brasil = calcular_indicadores_ano(
        bases_medica[2026],
        bases_multi[2026],
        uf=None,
        municipio=None,
        instituicoes=None,
    )

    # ========================================================
    # SÉRIES = UF
    # ========================================================

    serie_estado = calcular_serie(
        bases_medica,
        bases_multi,
        uf=uf,
        municipio=None,
        instituicoes=None,
    )

    serie_brasil = calcular_serie(
        bases_medica,
        bases_multi,
        uf=None,
        municipio=None,
        instituicoes=None,
    )

    validar_valores(
        indicadores_uf,
        serie_estado,
    )

    if not MODELO_BRIEFING.exists():
        raise FileNotFoundError(
            f"Modelo não encontrado:\n{MODELO_BRIEFING}"
        )

    doc = Document(MODELO_BRIEFING)

    # ========================================================
    # 1 - TEXTOS GERAIS = UF
    # ========================================================

    preencher_textos(
        doc,
        uf,
        municipio,
        indicadores_uf,
    )

    # ========================================================
    # 2 - SEÇÃO 3.2 = MUNICÍPIO
    # ========================================================

    if municipio:
        indicadores_municipio = obter_dados_municipio_2026(
            bases_medica,
            bases_multi,
            uf,
            municipio,
        )

        preencher_texto_secao_32(
            doc,
            uf,
            municipio,
            indicadores_municipio,
        )

    # ========================================================
    # 3 - GRANDES NÚMEROS = UF
    # ========================================================

    preencher_grandes_numeros(
        doc,
        indicadores_uf,
        uf,
    )

    # ========================================================
    # 4 - VARIAÇÕES = UF / BRASIL
    # ========================================================

    preencher_variacoes_residencias(
        doc,
        serie_brasil,
        serie_estado,
    )

    # ========================================================
    # 5 - COMPARATIVO = BRASIL X UF
    # ========================================================

    preencher_comparativo(
        doc,
        indicadores_brasil,
        indicadores_uf,
        uf,
    )

    # ========================================================
    # 6 - SÉRIE HISTÓRICA = UF
    # ========================================================

    preencher_serie_historica(
        doc,
        serie_brasil,
        serie_estado,
    )

    # ========================================================
    # 7 - TOP 5 = MUNICÍPIO
    # ========================================================

    if municipio:
        preencher_tabelas_top5(
            doc,
            bases_medica,
            bases_multi,
            uf,
            municipio,
        )

    # ========================================================
    # VALIDAÇÃO FINAL
    # ========================================================

    restantes = encontrar_placeholders(doc)

    if restantes:
        print("\nATENÇÃO: placeholders [nº] restantes:")

        for local, texto in restantes:
            print(
                f"  {local}: {texto[:250]}"
            )

        raise RuntimeError(
            "O documento ainda possui placeholders [nº]. "
            "A geração foi interrompida."
        )

    # ========================================================
    # NOME DO ARQUIVO
    # ========================================================

    nome_municipio = (
        normalizar_texto(municipio)
        if municipio
        else "TODOS"
    )

    nome = (
        f"Briefing_{UF_ATUAL}_{nome_municipio}.docx"
    )

    caminho = PASTA_SAIDA / nome

    doc.save(caminho)

    # ========================================================
    # GRÁFICOS
    # ========================================================

    atualizar_graficos(
        caminho,
        serie_brasil,
        serie_estado,
    )

    # ========================================================
    # RESUMO
    # ========================================================

    print("\n" + "=" * 70)
    print("✓ BRIEFING GERADO COM SUCESSO")
    print("=" * 70)

    print(f"\nArquivo: {caminho}")

    print("\n2026 - ESTADO / UF")
    print(
        f"Programas Médica: "
        f"{formatar_numero(indicadores_uf['programas_medica'])}"
    )
    print(
        f"Programas Multi:   "
        f"{formatar_numero(indicadores_uf['programas_multi'])}"
    )
    print(
        f"Programas Total:   "
        f"{formatar_numero(indicadores_uf['programas_total'])}"
    )
    print(
        f"Bolsas Médica:     "
        f"{formatar_numero(indicadores_uf['bolsas_medica'])}"
    )
    print(
        f"Bolsas Multi:      "
        f"{formatar_numero(indicadores_uf['bolsas_multi'])}"
    )
    print(
        f"Bolsas Total:      "
        f"{formatar_numero(indicadores_uf['bolsas_total'])}"
    )
    print(
        f"Financeiro Médica: "
        f"{formatar_moeda(indicadores_uf['financeiro_medica'])}"
    )
    print(
        f"Financeiro Multi:  "
        f"{formatar_moeda(indicadores_uf['financeiro_multi'])}"
    )
    print(
        f"Financeiro Total:  "
        f"{formatar_moeda(indicadores_uf['financeiro_total'])}"
    )

    if municipio:
        print(f"\n2026 - MUNICÍPIO: {municipio}")

        indicadores_municipio = obter_dados_municipio_2026(
            bases_medica,
            bases_multi,
            uf,
            municipio,
        )

        print(
            f"Programas Médica: "
            f"{formatar_numero(indicadores_municipio['programas_medica'])}"
        )
        print(
            f"Programas Multi:   "
            f"{formatar_numero(indicadores_municipio['programas_multi'])}"
        )
        print(
            f"Programas Total:   "
            f"{formatar_numero(indicadores_municipio['programas_total'])}"
        )
        print(
            f"Bolsas Médica:     "
            f"{formatar_numero(indicadores_municipio['bolsas_medica'])}"
        )
        print(
            f"Bolsas Multi:      "
            f"{formatar_numero(indicadores_municipio['bolsas_multi'])}"
        )
        print(
            f"Bolsas Total:      "
            f"{formatar_numero(indicadores_municipio['bolsas_total'])}"
        )

    print("\nSÉRIE DO ESTADO")
    for ano in ANOS:
        print(
            f"{ano}: "
            f"{formatar_numero(serie_estado[ano]['programas_total'])} programas | "
            f"{formatar_numero(serie_estado[ano]['bolsas_total'])} bolsas | "
            f"{formatar_moeda(serie_estado[ano]['financeiro_total'])}"
        )

    print("\nSÉRIE BRASIL")
    for ano in ANOS:
        print(
            f"{ano}: "
            f"{formatar_numero(serie_brasil[ano]['programas_total'])} programas | "
            f"{formatar_numero(serie_brasil[ano]['bolsas_total'])} bolsas | "
            f"{formatar_moeda(serie_brasil[ano]['financeiro_total'])}"
        )

    return caminho


# ============================================================
# TERMINAL
# ============================================================

def escolher_opcao(
    titulo,
    opcoes,
    permitir_todos=False,
):
    print("\n" + "=" * 70)
    print(titulo)
    print("=" * 70)

    if permitir_todos:
        print("0 - Todos")

    for i, opcao in enumerate(opcoes, start=1):
        print(f"{i} - {opcao}")

    while True:
        entrada = input(
            "\nDigite o número da opção: "
        ).strip()

        try:
            n = int(entrada)
        except ValueError:
            print("Digite somente um número.")
            continue

        if permitir_todos and n == 0:
            return None

        if 1 <= n <= len(opcoes):
            return opcoes[n - 1]

        print("Opção inválida.")


def escolher_instituicoes(instituicoes):
    print("\n" + "=" * 70)
    print("INSTITUIÇÕES")
    print("=" * 70)

    print("0 - Todas")

    for i, instituicao in enumerate(
        instituicoes,
        start=1,
    ):
        print(f"{i} - {instituicao}")

    while True:
        entrada = input(
            "\nDigite 0 ou números separados por vírgula: "
        ).strip()

        if entrada == "0":
            return []

        try:
            numeros = [
                int(x.strip())
                for x in entrada.split(",")
            ]
        except ValueError:
            print("Entrada inválida.")
            continue

        if not numeros:
            print("Selecione pelo menos uma instituição.")
            continue

        if any(
            n < 1 or n > len(instituicoes)
            for n in numeros
        ):
            print("Existe uma opção inválida.")
            continue

        return [
            instituicoes[n - 1]
            for n in numeros
        ]


def executar():
    print("=" * 70)
    print("GERADOR DE BRIEFING SGTES")
    print("=" * 70)

    bases_medica, bases_multi = carregar_bases()

    while True:

        ufs = listar_ufs(
            bases_medica,
            bases_multi,
        )

        uf = escolher_opcao(
            "SELECIONE A UF",
            ufs,
        )

        municipios = listar_municipios(
            bases_medica,
            bases_multi,
            uf,
        )

        municipio = escolher_opcao(
            "SELECIONE O MUNICÍPIO",
            municipios,
            permitir_todos=True,
        )

        instituicoes = listar_instituicoes(
            bases_medica,
            bases_multi,
            uf,
            municipio,
        )

        # Mantido para compatibilidade com o fluxo anterior.
        # A seleção de instituições NÃO interfere no Top 5.
        if instituicoes:
            selecionadas = escolher_instituicoes(
                instituicoes,
            )
        else:
            selecionadas = []

        print("\n" + "=" * 70)
        print("SELEÇÃO")
        print("=" * 70)

        print(f"UF: {uf}")
        print(
            f"Município: "
            f"{municipio or 'Todos os municípios'}"
        )

        print(
            f"Instituições selecionadas: "
            f"{len(selecionadas) if selecionadas else 'Todas'}"
        )

        confirmar = input(
            "\nGerar briefing? [S/N]: "
        ).strip().upper()

        if confirmar not in {"S", "SIM"}:
            continue

        try:
            gerar_briefing(
                bases_medica,
                bases_multi,
                uf,
                municipio,
                selecionadas,
            )

        except Exception as exc:
            print("\n" + "=" * 70)
            print("✗ ERRO AO GERAR O BRIEFING")
            print("=" * 70)
            print(
                f"\n{type(exc).__name__}: {exc}"
            )

        resposta = input(
            "\nDeseja gerar outro briefing? [S/N]: "
        ).strip().upper()

        if resposta not in {"S", "SIM"}:
            break


# ============================================================
# MAIN
# ============================================================

if __name__ == "__main__":
    executar()

    
